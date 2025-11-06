"""File system scanner for documents."""

import os
import hashlib
from pathlib import Path
from typing import List, Optional, Dict
from datetime import datetime

from app.config import settings
from app.database import Document, get_db_session


def calculate_md5(file_path: str) -> str:
    """Calculate MD5 hash of a file."""
    md5_hash = hashlib.md5()
    try:
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(settings.chunk_size), b""):
                md5_hash.update(chunk)
        return md5_hash.hexdigest()
    except (IOError, OSError) as e:
        raise IOError(f"Error calculating MD5 for {file_path}: {e}")


def get_file_metadata(file_path: str) -> Dict:
    """Extract file metadata."""
    path_obj = Path(file_path)
    stat_info = path_obj.stat()

    # Extract drive letter (Windows)
    drive = file_path[0] if file_path and file_path[1] == ':' else ""
    directory = str(path_obj.parent)

    return {
        "name": path_obj.stem,
        "file_path": file_path,
        "drive": drive,
        "directory": directory,
        "size": stat_info.st_size,
        "size_on_disc": stat_info.st_size,  # Simplified for now
        "date_created": datetime.fromtimestamp(stat_info.st_ctime),
        "date_published": None,  # Will extract from PDF metadata if available
        "file_type": path_obj.suffix.lower(),
    }


def scan_drive(drive_letter: str,
               file_extensions: Optional[List[str]] = None,
               max_files: Optional[int] = None) -> List[str]:
    """
    Scan a drive for documents.

    Args:
        drive_letter: Drive letter (e.g., 'C', 'D')
        file_extensions: List of file extensions to scan
        max_files: Maximum number of files to return (None = no limit)

    Returns:
        List of file paths found
    """
    if file_extensions is None:
        file_extensions = settings.supported_extensions

    root_path = f"{drive_letter}:\\"
    found_files = []

    if not os.path.exists(root_path):
        raise ValueError(f"Drive {root_path} does not exist")

    for root, dirs, files in os.walk(root_path):
        # Skip hidden directories
        dirs[:] = [d for d in dirs if not d.startswith('.')]

        for file in files:
            file_path = os.path.join(root, file)
            file_ext = os.path.splitext(file_path)[1].lower()

            if file_ext in file_extensions:
                found_files.append(file_path)
                # Stop early if max_files limit reached
                if max_files is not None and len(found_files) >= max_files:
                    return found_files

    return found_files


def scan_all_drives(file_extensions: Optional[List[str]] = None,
                    max_files: Optional[int] = None) -> Dict[str, List[str]]:
    """
    Scan all available drives.

    Args:
        file_extensions: List of file extensions to scan
        max_files: Maximum number of files per drive (None = no limit)

    Returns:
        Dictionary mapping drive letters to lists of file paths
    """
    drives_by_letter = {}

    # Get all drive letters (Windows)
    import string
    for drive_letter in string.ascii_uppercase:
        drive_path = f"{drive_letter}:\\"
        if os.path.exists(drive_path):
            try:
                files = scan_drive(drive_letter, file_extensions, max_files)
                if files:
                    drives_by_letter[drive_letter] = files
            except (PermissionError, OSError) as e:
                print(f"Error scanning drive {drive_letter}: {e}")

    return drives_by_letter


def index_document(file_path: str,
                   extract_text: bool = True) -> Optional[Document]:
    """
    Index a document and store in database.

    Args:
        file_path: Path to the document
        extract_text: Whether to extract text content

    Returns:
        Document object if successful, None otherwise
    """
    try:
        # Check if file exists and is readable
        if not os.path.exists(file_path) or not os.path.isfile(file_path):
            return None

        # Get metadata
        metadata = get_file_metadata(file_path)

        # Calculate MD5
        md5_hash = calculate_md5(file_path)

        # Check if already indexed
        from app.database import SessionLocal
        db = SessionLocal()
        try:
            # Clean up stale database records for files in the same directory that no longer exist
            # This ensures no mirage database records - reconcile automatically during scanning
            try:
                file_dir = os.path.dirname(file_path)
                if file_dir:
                    # Get all documents in the same directory
                    docs_in_dir = db.query(Document).filter(
                        Document.directory == file_dir
                    ).all()
                    
                    # Remove records for files that no longer exist
                    deleted_count = 0
                    for doc in docs_in_dir:
                        if not os.path.exists(doc.file_path):
                            try:
                                db.delete(doc)
                                deleted_count += 1
                            except Exception:
                                # Silently handle database errors - don't show to user
                                try:
                                    db.rollback()
                                except Exception:
                                    pass
                                continue
                    
                    # Commit cleanup if any deletions were made
                    if deleted_count > 0:
                        try:
                            db.commit()
                        except Exception:
                            # Silently handle database errors - don't show to user
                            try:
                                db.rollback()
                            except Exception:
                                pass
            except Exception:
                # Silently handle cleanup errors - don't fail indexing
                try:
                    db.rollback()
                except Exception:
                    pass
            
            existing = db.query(Document).filter(
                Document.file_path == file_path
            ).first()

            if existing:
                # Update if needed
                for key, value in metadata.items():
                    setattr(existing, key, value)
                existing.md5_hash = md5_hash
                if extract_text:
                    extracted_text = extract_text_content(file_path)
                    existing.extracted_text = extracted_text
                    # Update preview
                    if extracted_text:
                        existing.extracted_text_preview = extracted_text[:8192]
                db.commit()
                return existing

            # Extract text if requested
            extracted_text = None
            extracted_text_preview = None
            
            if extract_text:
                extracted_text = extract_text_content(file_path)
                # Store first 8KB for preview
                if extracted_text:
                    extracted_text_preview = extracted_text[:8192]

            # Create document record
            document = Document(
                name=metadata["name"],
                file_path=metadata["file_path"],
                drive=metadata["drive"],
                directory=metadata["directory"],
                size=metadata["size"],
                size_on_disc=metadata["size_on_disc"],
                date_created=metadata["date_created"],
                date_published=metadata["date_published"],
                md5_hash=md5_hash,
                file_type=metadata["file_type"],
                extracted_text=extracted_text,
                extracted_text_preview=extracted_text_preview,
            )

            # Extract author if available from PDF metadata
            # Wrap in try/except to handle corrupted PDFs gracefully
            if metadata["file_type"] == ".pdf":
                try:
                    author = extract_pdf_author(file_path)
                    if author:
                        document.author = author
                except Exception:
                    # Silently skip if author extraction fails
                    pass

            db.add(document)
            db.commit()
            db.refresh(document)
            return document
        finally:
            db.close()

    except Exception as e:
        print(f"Error indexing {file_path}: {e}")
        return None


def extract_text_content(file_path: str) -> Optional[str]:
    """Extract text content from document."""
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == ".pdf":
        return extract_pdf_text(file_path)
    elif file_ext == ".txt":
        return extract_txt_text(file_path)
    elif file_ext == ".docx":
        return extract_docx_text(file_path)
    elif file_ext == ".epub":
        return extract_epub_text(file_path)

    return None


def extract_pdf_text(file_path: str) -> Optional[str]:
    """Extract text from PDF file."""
    import warnings
    import io
    import sys
    
    # Suppress PDF warnings and errors
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        # Redirect stderr to suppress PDF library errors
        old_stderr = sys.stderr
        sys.stderr = io.StringIO()
        
        try:
            try:
                import pdfplumber
                text_content = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        try:
                            text = page.extract_text()
                            if text:
                                text_content.append(text)
                        except Exception:
                            # Skip corrupted pages, continue with next page
                            continue
                result = "\n".join(text_content) if text_content else None
                sys.stderr = old_stderr
                return result
            except Exception:
                # Try PyPDF2 as fallback for corrupted PDFs
                try:
                    import PyPDF2
                    text_content = []
                    with open(file_path, "rb") as f:
                        pdf_reader = PyPDF2.PdfReader(f, strict=False)
                        for page in pdf_reader.pages:
                            try:
                                text = page.extract_text()
                                if text:
                                    text_content.append(text)
                            except Exception:
                                continue
                    result = "\n".join(text_content) if text_content else None
                    sys.stderr = old_stderr
                    return result
                except Exception:
                    # Silently skip if both methods fail
                    sys.stderr = old_stderr
                    return None
        except Exception:
            sys.stderr = old_stderr
            return None


def extract_pdf_author(file_path: str) -> Optional[str]:
    """Extract author metadata from PDF."""
    import warnings
    import io
    import sys
    
    # Suppress PDF warnings and errors
    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        # Redirect stderr to suppress PDF library errors
        old_stderr = sys.stderr
        sys.stderr = io.StringIO()
        
        try:
            import PyPDF2
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f, strict=False)
                if pdf_reader.metadata and pdf_reader.metadata.author:
                    author = pdf_reader.metadata.author
                    sys.stderr = old_stderr
                    return author
            sys.stderr = old_stderr
            return None
        except Exception:
            # Silently skip corrupted PDFs - don't print error to avoid spam
            sys.stderr = old_stderr
            return None


def extract_txt_text(file_path: str) -> Optional[str]:
    """Extract text from TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"Error extracting TXT text from {file_path}: {e}")
        return None


def extract_docx_text(file_path: str) -> Optional[str]:
    """Extract text from DOCX file."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error extracting DOCX text from {file_path}: {e}")
        return None


def extract_epub_text(file_path: str) -> Optional[str]:
    """Extract text from EPUB file."""
    try:
        import ebooklib
        from ebooklib import epub
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            # Fallback if BeautifulSoup not available
            from html.parser import HTMLParser

            class TextExtractor(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text = []

                def handle_data(self, data):
                    self.text.append(data)

                def get_text(self):
                    return " ".join(self.text)

            soup_class = TextExtractor

        book = epub.read_epub(file_path)
        text_content = []

        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                content = item.get_content()
                try:
                    soup = BeautifulSoup(content, 'html.parser')
                    text_content.append(soup.get_text())
                except NameError:
                    # Use fallback parser
                    parser = TextExtractor()
                    parser.feed(content.decode('utf-8', errors='ignore'))
                    text_content.append(parser.get_text())

        return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting EPUB text from {file_path}: {e}")
        return None


def find_duplicates() -> Dict[str, List[Document]]:
    """
    Find duplicate documents by MD5 hash (same content, different names).

    Returns:
        Dictionary mapping MD5 hash to list of duplicate documents
    """
    from app.database import SessionLocal
    db = SessionLocal()
    try:
        documents = db.query(Document).all()

        hash_groups = {}
        for doc in documents:
            if doc.md5_hash not in hash_groups:
                hash_groups[doc.md5_hash] = []
            hash_groups[doc.md5_hash].append(doc)

        # Return only groups with duplicates
        duplicates = {h: docs for h, docs in hash_groups.items()
                      if len(docs) > 1}

        return duplicates
    finally:
        db.close()


def find_duplicate_by_name() -> Dict[str, List[Document]]:
    """
    Find duplicate documents by name (same name, different content).

    Returns:
        Dictionary mapping normalized name to list of documents with same name
    """
    from app.database import SessionLocal
    db = SessionLocal()
    try:
        documents = db.query(Document).all()

        name_groups = {}
        for doc in documents:
            # Normalize name (lowercase, strip whitespace)
            normalized_name = doc.name.lower().strip()
            if normalized_name not in name_groups:
                name_groups[normalized_name] = []
            name_groups[normalized_name].append(doc)

        # Return only groups with same name but different MD5 (different content)
        duplicates = {}
        for name, docs in name_groups.items():
            if len(docs) > 1:
                # Check if they have different MD5 hashes
                md5_set = set(doc.md5_hash for doc in docs)
                if len(md5_set) > 1:
                    # Same name, different content
                    duplicates[name] = docs

        return duplicates
    finally:
        db.close()


def find_all_duplicates() -> Dict:
    """
    Find all types of duplicates:
    1. Same name, different content (different MD5)
    2. Different name, same content (same MD5)

    Returns:
        Dictionary with two types of duplicates
    """
    same_content_diff_name = find_duplicates()  # Same MD5, different names
    same_name_diff_content = find_duplicate_by_name()  # Same name, different MD5

    return {
        "same_content_diff_name": same_content_diff_name,
        "same_name_diff_content": same_name_diff_content,
        "total_same_content": sum(len(docs) - 1 for docs in same_content_diff_name.values()),
        "total_same_name": sum(len(docs) - 1 for docs in same_name_diff_content.values()),
    }


def calculate_space_savings(duplicates: Dict[str, List[Document]],
                            keep_location: str) -> int:
    """
    Calculate space savings from removing duplicates.

    Args:
        duplicates: Dictionary of duplicate groups
        keep_location: Drive or directory to keep files in

    Returns:
        Total space that would be saved in bytes
    """
    total_savings = 0

    for hash_val, docs in duplicates.items():
        # Find document to keep (prefer specified location)
        keep_doc = None
        for doc in docs:
            if keep_location.lower() in doc.directory.lower():
                keep_doc = doc
                break

        if not keep_doc:
            keep_doc = docs[0]

        # Calculate savings from other duplicates
        for doc in docs:
            if doc.id != keep_doc.id:
                total_savings += doc.size

    return total_savings


def delete_duplicate_file(file_path: str,
                         space_saved: int,
                         user_id: Optional[int] = None) -> None:
    """
    Delete a duplicate file and log the activity.

    Args:
        file_path: Path to the file to delete
        space_saved: Space saved by deletion
        user_id: User ID who performed the deletion
    """
    import os
    from app.reports import log_activity

    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            log_activity(
                activity_type="delete",
                description=f"Deleted duplicate file: {file_path}",
                document_path=file_path,
                space_saved_bytes=space_saved,
                operation_count=1,
                user_id=user_id
            )
    except Exception as e:
        raise Exception(f"Error deleting {file_path}: {e}")


def move_document(source_path: str,
                  target_path: str,
                  user_id: Optional[int] = None) -> None:
    """
    Move a document and log the activity.

    Args:
        source_path: Source file path
        target_path: Target file path
        user_id: User ID who performed the move
    """
    import shutil
    from app.reports import log_activity

    try:
        shutil.move(source_path, target_path)
        log_activity(
            activity_type="move",
            description=f"Moved document from {source_path} to {target_path}",
            document_path=target_path,
            space_saved_bytes=0,
            operation_count=1,
            user_id=user_id
        )
    except Exception as e:
        raise Exception(f"Error moving {source_path}: {e}")


